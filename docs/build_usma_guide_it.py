"""Build the Italian USMA user guide and theory wiki DOCX.

The document is intentionally generated from a script so the guide can be
rebuilt when the application, theory notes, or example logs change.
"""

from __future__ import annotations

import math
import re
import shutil
from pathlib import Path
from typing import Iterable, Sequence

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "usma_guide_assets"
DOCX_PATH = OUT_DIR / "USMA_Guida_Utente_e_Wiki_IT.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "162033"
MUTED = "5B6472"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CFD6E0"
GREEN = "2EAD68"
ORANGE = "E67E22"
RED = "C0392B"
TEAL = "1ABC9C"
PURPLE = "7E57C2"
YELLOW = "F1C40F"


def get_version() -> str:
    text = (ROOT / "usma" / "models.py").read_text(encoding="utf-8")
    match = re.search(r'APP_VERSION\s*=\s*"([^"]+)"', text)
    return match.group(1) if match else "0.12.4"


APP_VERSION = get_version()


def ensure_dirs() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: str = "#162033",
    line_gap: int = 4,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 22)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) // 2 - 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str = "#CFD6E0",
    radius: int = 18,
    width: int = 2,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#5B6472") -> None:
    draw.line([start, end], fill=fill, width=4)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    pts = [
        end,
        (end[0] - size * math.cos(ang - math.pi / 7), end[1] - size * math.sin(ang - math.pi / 7)),
        (end[0] - size * math.cos(ang + math.pi / 7), end[1] - size * math.sin(ang + math.pi / 7)),
    ]
    draw.polygon(pts, fill=fill)


def save_pipeline_diagram(path: Path) -> None:
    img = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(36, True)
    h2 = font(24, True)
    body = font(21)
    draw.text((62, 40), "Flusso dati USMA", font=h1, fill=f"#{INK}")
    draw.text((62, 88), "Dalla schermata al verdetto e ai file esportati", font=body, fill=f"#{MUTED}")

    boxes = [
        ((70, 185, 285, 285), "Cattura schermo", "#EAF3FF"),
        ((360, 185, 575, 285), "ROI definite", "#F4F6F9"),
        ((650, 185, 865, 285), "Filtro HSV", "#FFF7E6"),
        ((940, 185, 1155, 285), "Segnale 1-D", "#EAFBF6"),
        ((1230, 185, 1445, 285), "Analisi", "#F0ECFF"),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill)
        text_center(draw, box, label, h2)
    for i in range(len(boxes) - 1):
        arrow(draw, (boxes[i][0][2] + 16, 235), (boxes[i + 1][0][0] - 16, 235))

    branches = [
        ((1000, 400, 1215, 485), "FRF / PSD\nFFT + Lowpass", "#EAF3FF"),
        ((1000, 520, 1215, 605), "Coerenza\nbadness + trend", "#EAFBF6"),
        ((1265, 400, 1480, 485), "OCR\nstatus, run, punti", "#F4F6F9"),
        ((1265, 520, 1480, 605), "Output\nlog immagini + UNV", "#FFF7E6"),
    ]
    for box, label, fill in branches:
        rounded_box(draw, box, fill)
        text_center(draw, box, label, body)
    arrow(draw, (1338, 292), (1106, 392))
    arrow(draw, (1338, 292), (1106, 512))
    arrow(draw, (1338, 292), (1372, 392))
    arrow(draw, (1106, 492), (1372, 512))
    arrow(draw, (1372, 492), (1372, 512))

    rounded_box(draw, (70, 405, 820, 605), "#F8FAFC")
    draw.text((105, 430), "Punto chiave", font=h2, fill=f"#{DARK_BLUE}")
    msg = (
        "Le decisioni FFT e Lowpass lavorano sul segnale in pixel ricostruito dal filtro HSV. "
        "La conversione in unità fisiche serve per lettura, diagnostica ed esportazione UNV."
    )
    y = 475
    for line in wrap(draw, msg, body, 670):
        draw.text((105, y), line, font=body, fill=f"#{INK}")
        y += 30
    img.save(path)


def save_classification_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(34, True)
    h2 = font(24, True)
    body = font(20)
    small = font(18)
    draw.text((60, 45), "Logica di classificazione", font=h1, fill=f"#{INK}")
    draw.text((60, 92), "Coerenza esclusa dal verdetto GOOD / SUSPECT / BAD", font=body, fill=f"#{MUTED}")

    # Per-signal matrix.
    x0, y0 = 115, 190
    cell_w, cell_h = 255, 120
    draw.text((x0 + cell_w + 20, y0 - 55), "Lowpass (colonne)", font=h2, fill=f"#{DARK_BLUE}")
    draw.text((x0, y0 + 378), "FFT (righe)", font=h2, fill=f"#{DARK_BLUE}")
    headers = ["OK", "BAD"]
    for i, label in enumerate(headers):
        rounded_box(draw, (x0 + (i + 1) * cell_w, y0, x0 + (i + 2) * cell_w, y0 + cell_h), "#F2F4F7")
        text_center(draw, (x0 + (i + 1) * cell_w, y0, x0 + (i + 2) * cell_w, y0 + cell_h), label, h2)
        rounded_box(draw, (x0, y0 + (i + 1) * cell_h, x0 + cell_w, y0 + (i + 2) * cell_h), "#F2F4F7")
        text_center(draw, (x0, y0 + (i + 1) * cell_h, x0 + cell_w, y0 + (i + 2) * cell_h), label, h2)
    matrix = [
        ("GOOD HIT", GREEN),
        ("SUSPECT\n(Lowpass)", ORANGE),
        ("SUSPECT\n(FFT)", ORANGE),
        ("BAD HIT", RED),
    ]
    for r in range(2):
        for c in range(2):
            idx = r * 2 + c
            fill = f"#{matrix[idx][1]}"
            box = (x0 + (c + 1) * cell_w, y0 + (r + 1) * cell_h, x0 + (c + 2) * cell_w, y0 + (r + 2) * cell_h)
            rounded_box(draw, box, fill, outline="#FFFFFF")
            text_center(draw, box, matrix[idx][0], h2, "white")

    # Cross FRF/PSD rule.
    rounded_box(draw, (875, 190, 1435, 665), "#F8FAFC")
    draw.text((915, 225), "Regola globale", font=h2, fill=f"#{DARK_BLUE}")
    rules = [
        ("Nessun metodo fallisce", "GOOD HIT", GREEN),
        ("Fallisce un solo lato o un solo metodo", "SUSPECT", ORANGE),
        ("FRF e PSD hanno almeno un voto BAD ciascuno", "BAD HIT", RED),
        ("Se esiste solo FRF o solo PSD: servono FFT e LP BAD", "BAD HIT", RED),
    ]
    y = 285
    for left, right, color in rules:
        draw.ellipse((915, y + 6, 937, y + 28), fill=f"#{color}")
        draw.text((955, y), left, font=small, fill=f"#{INK}")
        draw.text((955, y + 28), right, font=body, fill=f"#{color}")
        y += 88
    draw.text((915, 615), "La coerenza resta un indicatore di qualità di misura.", font=small, fill=f"#{MUTED}")

    # Method lights.
    labels = [("FRF-FFT", BLUE), ("FRF-LP", TEAL), ("PSD-FFT", PURPLE), ("PSD-LP", ORANGE)]
    x = 140
    y = 710
    for label, color in labels:
        draw.ellipse((x, y, x + 54, y + 54), fill=f"#{color}")
        draw.text((x + 70, y + 12), label, font=body, fill=f"#{INK}")
        x += 300
    img.save(path)


def save_calibration_levels(path: Path) -> None:
    img = Image.new("RGB", (1500, 620), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(34, True)
    h2 = font(22, True)
    body = font(19)
    draw.text((60, 45), "Scala di confidenza della calibrazione", font=h1, fill=f"#{INK}")
    draw.text((60, 92), "Livelli ordinali: indicano quantità di dati e robustezza del merge, non una probabilità assoluta.", font=body, fill=f"#{MUTED}")

    levels = [
        ("0", "< 6", "Default", RED, "Non calibrato"),
        ("1", "6-7", "Percentili", ORANGE, "Stima preliminare"),
        ("2", "8-11", "Percentili + Bayes", YELLOW, "Base"),
        ("3", "12-15", "Tutti e tre", BLUE, "Cross-validato"),
        ("4", "16+", "Tutti e tre", GREEN, "Robusto"),
    ]
    x = 80
    y = 205
    w = 260
    for lvl, n, method, color, meaning in levels:
        rounded_box(draw, (x, y, x + 230, y + 260), "#F8FAFC", outline=f"#{color}", width=4)
        draw.ellipse((x + 83, y + 25, x + 147, y + 89), fill=f"#{color}")
        text_center(draw, (x + 83, y + 25, x + 147, y + 89), lvl, h2, "white")
        text_center(draw, (x + 15, y + 105, x + 215, y + 150), f"{n} segnali", h2)
        text_center(draw, (x + 15, y + 158, x + 215, y + 202), method, body)
        text_center(draw, (x + 15, y + 208, x + 215, y + 247), meaning, body, f"#{MUTED}")
        x += w
    img.save(path)


def save_hsv_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 700), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(34, True)
    h2 = font(24, True)
    body = font(20)
    mono = font(19)
    draw.text((60, 45), "Filtro HSV", font=h1, fill=f"#{INK}")
    draw.text((60, 92), "Un pixel entra nella maschera solo se rispetta tutti e tre gli intervalli.", font=body, fill=f"#{MUTED}")

    channels = [
        ("H", "Hue", "0 - 179", "#E74C3C"),
        ("S", "Saturazione", "0 - 255", "#3498DB"),
        ("V", "Valore", "0 - 255", "#F1C40F"),
    ]
    y = 185
    for letter, name, rng, color in channels:
        rounded_box(draw, (90, y, 1320, y + 82), "#F8FAFC")
        draw.ellipse((125, y + 17, 173, y + 65), fill=color)
        text_center(draw, (125, y + 17, 173, y + 65), letter, h2, "white")
        draw.text((205, y + 17), name, font=h2, fill=f"#{INK}")
        draw.text((480, y + 22), rng, font=body, fill=f"#{MUTED}")
        draw.rounded_rectangle((720, y + 28, 1260, y + 56), radius=12, fill="#E5E7EB")
        # Example defaults are broad; upper V excludes pure white.
        if letter == "V":
            draw.rounded_rectangle((720, y + 28, 1228, y + 56), radius=12, fill=color)
            draw.text((1268, y + 25), "upper 240", font=mono, fill=f"#{MUTED}")
        else:
            draw.rounded_rectangle((720, y + 28, 1260, y + 56), radius=12, fill=color)
        y += 115

    rounded_box(draw, (90, 535, 1320, 635), "#FFF7E6", outline="#F3D19C")
    formula = "mask = (H_min <= H <= H_max) AND (S_min <= S <= S_max) AND (V_min <= V <= V_max)"
    text_center(draw, (110, 535, 1300, 635), formula, mono)
    img.save(path)


def save_reconstruction_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(34, True)
    h2 = font(23, True)
    body = font(20)
    mono = font(19)
    draw.text((60, 45), "Ricostruzione del segnale", font=h1, fill=f"#{INK}")
    draw.text((60, 92), "La maschera HSV viene compressa in un segnale 1-D colonna per colonna.", font=body, fill=f"#{MUTED}")

    plot_box = (100, 185, 720, 600)
    draw.rectangle(plot_box, outline=f"#{BORDER}", width=2, fill="#FCFCFD")
    x_vals = np.linspace(0, 1, 90)
    y_vals = 0.45 + 0.23 * np.sin(2 * np.pi * x_vals * 1.8) + 0.08 * np.sin(2 * np.pi * x_vals * 6)
    pts = []
    for i, val in enumerate(y_vals):
        x = plot_box[0] + 35 + i * ((plot_box[2] - plot_box[0] - 70) / (len(y_vals) - 1))
        y = plot_box[3] - 40 - val * (plot_box[3] - plot_box[1] - 90)
        pts.append((x, y))
        draw.ellipse((x - 3, y - 3, x + 3, y + 3), fill="#00B8D9")
    draw.line(pts, fill="#00B8D9", width=4)
    for i in range(0, len(pts), 8):
        x, y = pts[i]
        draw.line((x, plot_box[3] - 40, x, y), fill="#D7DEE8", width=1)
    draw.text((140, 625), "mean row per colonna -> interpolazione -> inversione Y", font=body, fill=f"#{INK}")

    rounded_box(draw, (815, 205, 1390, 335), "#EAFBF6")
    text_center(draw, (835, 205, 1370, 335), "signal_pixels = height - mean_row", mono)
    rounded_box(draw, (815, 380, 1390, 560), "#EAF3FF")
    text_center(
        draw,
        (835, 380, 1370, 560),
        "signal_physical[i] = y_min + (signal_pixels[i] / height) * (y_max - y_min)",
        mono,
    )
    rounded_box(draw, (100, 690, 1390, 760), "#F8FAFC")
    text_center(draw, (125, 690, 1365, 760), "FFT e Lowpass usano signal_pixels; signal_physical serve per grafici ed export UNV.", body)
    img.save(path)


def save_coherence_diagram(path: Path) -> None:
    img = Image.new("RGB", (1500, 780), "white")
    draw = ImageDraw.Draw(img)
    h1 = font(34, True)
    h2 = font(23, True)
    body = font(20)
    mono = font(19)
    draw.text((60, 45), "Coerenza: badness normalizzata", font=h1, fill=f"#{INK}")
    draw.text((60, 92), "USMA integra l'area di 1 - gamma^2 sulla banda visibile.", font=body, fill=f"#{MUTED}")

    left, top, right, bottom = 110, 185, 1070, 620
    draw.rectangle((left, top, right, bottom), fill="#FCFCFD", outline=f"#{BORDER}", width=2)
    for i in range(6):
        y = top + i * (bottom - top) / 5
        draw.line((left, y, right, y), fill="#E5E7EB")
    for i in range(8):
        x = left + i * (right - left) / 7
        draw.line((x, top, x, bottom), fill="#EEF1F5")
    xs = np.linspace(0, 1, 160)
    gamma = 0.92 - 0.11 * np.exp(-((xs - 0.35) / 0.08) ** 2) - 0.22 * np.exp(-((xs - 0.72) / 0.07) ** 2)
    gamma += 0.025 * np.sin(2 * np.pi * xs * 3)
    gamma = np.clip(gamma, 0, 1)
    pts = []
    area_pts = [(left, top)]
    for xval, yval in zip(xs, gamma):
        x = left + xval * (right - left)
        y = bottom - yval * (bottom - top)
        pts.append((x, y))
        area_pts.append((x, y))
    area_pts.append((right, top))
    draw.polygon(area_pts, fill="#FFE6E2")
    draw.line(pts, fill=f"#{RED}", width=4)
    draw.text((left + 22, top + 22), "area = 1 - gamma^2", font=body, fill=f"#{RED}")
    draw.text((left, bottom + 18), "frequenza", font=body, fill=f"#{MUTED}")
    draw.text((left - 58, top + 10), "gamma^2", font=body, fill=f"#{MUTED}")

    rounded_box(draw, (1130, 225, 1400, 365), "#FFF7E6")
    text_center(draw, (1150, 225, 1380, 365), "B = integral(1 - gamma^2) df / delta_f", mono)
    rounded_box(draw, (1130, 415, 1400, 555), "#F8FAFC")
    text_center(draw, (1150, 415, 1380, 555), "Flag default: B > 0.30 oppure degrado relativo > 20%", body)
    img.save(path)


def copy_real_examples() -> dict[str, Path]:
    mapping = {
        "real_signal": ROOT / "image_logs" / "Signals" / "region_1_P1P1_100_signal.png",
        "real_fft": ROOT / "image_logs" / "FFT" / "region_1_P1P1_100_fft.png",
        "real_lowpass": ROOT / "image_logs" / "Lowpass" / "region_1_P1P1_100_lowpass.png",
        "real_residual": ROOT / "image_logs" / "Residual" / "region_1_P1P1_100_residual.png",
        "real_roi": ROOT / "image_logs" / "ROIs" / "region_1_P1P1_100_region_1.jpg",
    }
    copied: dict[str, Path] = {}
    for key, src in mapping.items():
        if src.exists():
            suffix = src.suffix.lower()
            dst = ASSET_DIR / f"{key}{suffix}"
            shutil.copyfile(src, dst)
            copied[key] = dst
    return copied


def generate_assets() -> dict[str, Path]:
    ensure_dirs()
    assets = {
        "pipeline": ASSET_DIR / "pipeline.png",
        "classification": ASSET_DIR / "classification.png",
        "calibration_levels": ASSET_DIR / "calibration_levels.png",
        "hsv": ASSET_DIR / "hsv.png",
        "reconstruction": ASSET_DIR / "signal_reconstruction.png",
        "coherence": ASSET_DIR / "coherence_badness.png",
    }
    save_pipeline_diagram(assets["pipeline"])
    save_classification_diagram(assets["classification"])
    save_calibration_levels(assets["calibration_levels"])
    save_hsv_diagram(assets["hsv"])
    save_reconstruction_diagram(assets["reconstruction"])
    save_coherence_diagram(assets["coherence"])
    assets.update(copy_real_examples())
    return assets


def rgb(hex_value: str) -> RGBColor:
    hex_value = hex_value.strip("#")
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def paragraph_border_bottom(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: float = 9.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths:
        for col_idx, width in enumerate(widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(width)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=DARK_BLUE, size=9.4)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=9.2)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph()


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Formula")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9.4, color=INK)


def add_h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Formula" not in styles:
        styles.add_style("Formula", 1)
    formula = styles["Formula"]
    formula.font.name = "Consolas"
    formula._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    formula.font.size = Pt(9.4)
    formula.font.color.rgb = rgb(INK)
    formula.paragraph_format.left_indent = Inches(0.18)
    formula.paragraph_format.space_before = Pt(3)
    formula.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(f"USMA v{APP_VERSION} | Guida utente e wiki teorica")
    set_run_font(r, size=9, color=MUTED)
    paragraph_border_bottom(header, color=BORDER, size="4")

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc: Document, assets: dict[str, Path]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("USMA")
    set_run_font(r, size=34, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Guida utente e Wiki teorica")
    set_run_font(r, size=21, color=BLUE, bold=True)

    add_para(
        doc,
        "Manuale in italiano per configurare, calibrare e usare Unified Screen Monitoring Application "
        "nei flussi di prova modale con FRF, PSD, coerenza, OCR ed esportazione UNV Dataset 58.",
    )
    add_table(
        doc,
        ["Voce", "Dettaglio"],
        [
            ("Versione USMA", f"v{APP_VERSION}"),
            ("Destinatari", "Utenti che non conoscono USMA, studenti e operatori di laboratorio"),
            ("Fonti tecniche", "README.md, THEORY.md, usma/theory/pages, moduli Python applicativi"),
            ("Formato", "Guida operativa + appendice wiki teorica integrata"),
        ],
        widths=[1.65, 4.65],
    )
    add_image(doc, assets["pipeline"], "Figura 1 - Flusso dati USMA ricostruito da codice e documentazione.", width=6.25)
    add_callout(
        doc,
        "Promessa della guida",
        "Le immagini e i grafici inclusi sono generati da regole del codice o ripresi dai log reali presenti nel repository. "
        "Quando una figura è esemplificativa, la didascalia lo dichiara; quando è un log reale, viene indicata come tale.",
        fill="F8FAFC",
    )
    doc.add_page_break()


def add_index(doc: Document) -> None:
    add_h(doc, "Indice dei contenuti", 1)
    add_bullets(
        doc,
        [
            "Parte 1 - Guida operativa: avvio, configurazioni, ROI, HSV, calibrazione, monitoraggio, log ed export.",
            "Parte 2 - Wiki teorica: ricostruzione segnale, FFT, Lowpass, coerenza, classificazione, calibrazione, HSV e grafici diagnostici.",
            "Parte 3 - Tabelle di riferimento: parametri di default, troubleshooting e fonti tecniche.",
        ],
    )
    add_callout(
        doc,
        "Come leggere il documento",
        "Se devi usare USMA subito, leggi la Parte 1. Se devi capire perché un parametro cambia un verdetto, usa la Parte 2. "
        "Se devi controllare valori numerici o range, vai alle tabelle di riferimento.",
    )
    doc.add_page_break()


def add_operational_guide(doc: Document, assets: dict[str, Path]) -> None:
    add_h(doc, "Parte 1 - Guida operativa", 1)
    add_h(doc, "Cos'è USMA", 2)
    add_para(
        doc,
        "USMA, Unified Screen Monitoring Application, monitora in tempo reale una schermata tecnica, estrae regioni di interesse "
        "definite dall'utente, ricostruisce segnali FRF/PSD/coerenza da tracce colorate, legge metadati tramite OCR e produce "
        "diagnostica più esportazione UNV Dataset 58 per flussi di analisi modale.",
    )
    add_bullets(
        doc,
        [
            "È portabile: Python e Tesseract OCR sono inclusi nella distribuzione.",
            "Non richiede diritti amministrativi.",
            "Lavora su Windows 10+ a 64 bit.",
            "Usa configurazioni JSON nella cartella configs e salva log in logs, image_logs e signal_logs.",
        ],
    )

    add_h(doc, "Avvio rapido", 2)
    add_numbers(
        doc,
        [
            "Estrai tutti i file mantenendo la struttura delle cartelle.",
            "Avvia RUN_USMA_PORTABLE.bat.",
            "Carica una configurazione esistente oppure crea una nuova definizione ROI.",
            "Disegna le ROI sullo screenshot, scegli il tipo corretto e imposta le scale fisiche per FRF, PSD o coerenza.",
            "Apri Define Color Filter e regola il filtro HSV finché la maschera segue solo la traccia utile.",
            "Salva la configurazione JSON.",
            "Alla scelta parametri, usa i valori salvati/default oppure avvia Calibrate with Expert Feedback se il setup è nuovo.",
            "Premi Start Monitoring e interpreta il banner Overall insieme ai dettagli FRF/PSD/coerenza.",
        ],
    )

    add_h(doc, "Struttura cartelle", 2)
    add_table(
        doc,
        ["Percorso", "Uso"],
        [
            ("RUN_USMA_PORTABLE.bat", "Punto di avvio consigliato per l'utente."),
            ("monitor_app.py", "Launcher leggero che delega al package usma."),
            ("usma/", "Codice applicativo modulare: modelli, monitor, analisi, export e GUI."),
            ("configs/", "Configurazioni JSON con ROI, HSV, parametri e dati di calibrazione."),
            ("image_logs/", "PNG/JPG diagnostici: ROI, segnali, FFT, Lowpass, residui, riepiloghi, OCR."),
            ("signal_logs/", "File .unv esportati in Dataset 58."),
            ("external/tesseract/", "Tesseract OCR portabile."),
            ("python/", "Python portabile usato dall'applicazione."),
        ],
        widths=[2.0, 4.3],
    )

    add_h(doc, "Schermata principale", 2)
    add_table(
        doc,
        ["Area", "Cosa fa"],
        [
            ("Configuration", "Load carica un JSON; Edit apre lo strumento di definizione ROI; View modifica dimensioni/scaling finestra."),
            ("Controls", "Start/Stop Monitoring, Overlay, frequenza campionamento, modalità Events Only e Audio."),
            ("Manual POI Entry", "Run, Hammer e Response manuali quando non sono disponibili ROI OCR dedicate."),
            ("Logging", "Verbose Log, Image Logs e opzione .unv."),
            ("Classification Methods", "Luci per FRF-FFT, FRF-LP, PSD-FFT e PSD-LP; servono a leggere quale metodo ha votato BAD."),
            ("Analysis Parameters (Live)", "Parametri FRF e PSD modificabili a caldo, più Live Calibration."),
            ("Graph Viewer", "Naviga hit, tipo di grafico, sorgente segnale e diagnostica di calibrazione."),
            ("Console Output", "Pannello opzionale con log operativi."),
        ],
        widths=[1.9, 4.4],
    )

    add_h(doc, "Tipi di ROI", 2)
    add_table(
        doc,
        ["Tipo", "Scopo operativo"],
        [
            ("frf", "Cattura segnale FRF; classificazione FFT + Lowpass; esportabile in UNV."),
            ("psd", "Cattura segnale PSD con set parametri indipendente da FRF."),
            ("coherence", "Monitora coerenza e badness; non contribuisce al verdetto GOOD/SUSPECT/BAD."),
            ("averages", "Legge via OCR il numero di medie/averages."),
            ("status", "Legge stato del sistema, ad esempio Waiting, Measuring o Ready."),
            ("overload", "Legge indicatore di overload."),
            ("run", "Legge via OCR il numero di run."),
            ("hammer", "Legge punto e direzione del martello."),
            ("response", "Legge punto e direzione della risposta."),
        ],
        widths=[1.15, 5.15],
    )

    add_h(doc, "Creare o modificare una configurazione", 2)
    add_numbers(
        doc,
        [
            "Apri Edit dalla schermata principale oppure Create New ROI Definition all'avvio.",
            "Premi Take Screenshot: USMA cattura il monitor selezionato nello spinbox Monitor.",
            "Trascina un rettangolo sulla preview per creare una ROI; scegli il tipo dal dialog.",
            "Nel Region Editor rinomina la ROI, controlla x, y, width, height e lascia Enabled attivo.",
            "Per frf, psd e coherence imposta X-Min, X-Max, Y-Min, Y-Max, Unit e Scale. Per coerenza usare normalmente Y-Min=0 e Y-Max=1.",
            "Scegli un colore overlay solo se aiuta a distinguere molte ROI; altrimenti lascia il default per tipo.",
            "Aggiorna la regione con Update, poi Save Config.",
        ],
    )
    add_callout(
        doc,
        "Nota sulle scale",
        "La classificazione FFT/Lowpass usa il segnale in pixel, quindi è robusta a errori di scala Y. "
        "La scala fisica resta però essenziale per grafici leggibili ed export UNV corretto.",
        fill="EAF3FF",
    )

    add_h(doc, "Calibrazione HSV", 2)
    add_image(doc, assets["hsv"], "Figura 2 - Regola esatta del filtro HSV usato da OpenCV in USMA.", width=6.25)
    add_bullets(
        doc,
        [
            "Apri Define Color Filter solo dopo aver definito almeno una ROI frf, psd o coherence.",
            "Scegli una ROI rappresentativa e regola H, S e V con anteprima live.",
            "La maschera è corretta quando la traccia è continua, senza griglia o sfondo e senza rumore sparso dominante.",
            "I default sono hsv_lower=[0,0,0] e hsv_upper=[179,255,240]; il valore V massimo a 240 esclude il bianco puro di sfondo/griglia.",
            "Premi Apply e poi salva la configurazione: cambiare gli slider senza salvare non rende persistente la modifica nel JSON.",
        ],
    )

    add_h(doc, "Calibrazione con feedback esperto", 2)
    add_image(doc, assets["calibration_levels"], "Figura 3 - Livelli di confidenza usati dal motore di calibrazione ibrido.", width=6.25)
    add_numbers(
        doc,
        [
            "Carica una configurazione e scegli Calibrate with Expert Feedback quando il setup è nuovo o i parametri non sono affidabili.",
            "Durante gli hit, classifica i segnali come Good, Bad o Skip. Se solo FRF o solo PSD va etichettato, usa i pulsanti specifici.",
            "Raccogli almeno 3 GOOD e 3 BAD per ogni famiglia di segnale da calibrare.",
            "Punta a 12 segnali per una calibrazione solida e a 16+ segnali per Level 4.",
            "Se USMA segnala un segnale simile, verifica se è davvero una ripetizione poco informativa; puoi comunque accettarlo se serve.",
            "Premi Finish Calibration per passare al monitoraggio normale con parametri stimati e salvati.",
        ],
    )
    add_callout(
        doc,
        "Persistenza",
        "I dati vengono salvati nel JSON sotto la chiave _calibration. Questo permette di riaprire la configurazione, continuare la calibrazione "
        "e ricalcolare le soglie con versioni successive dell'algoritmo.",
    )

    add_h(doc, "Monitoraggio e lettura dei risultati", 2)
    add_image(doc, assets["classification"], "Figura 4 - Logica di voto e aggregazione FRF/PSD implementata in usma/analysis/classifier.py.", width=6.25)
    add_table(
        doc,
        ["Indicazione", "Interpretazione"],
        [
            ("GOOD HIT", "Nessun metodo abilitato ha votato BAD."),
            ("SUSPECT", "Almeno un metodo ha votato BAD, ma non c'è accordo sufficiente per un BAD globale."),
            ("BAD HIT", "Con FRF e PSD presenti: almeno un metodo BAD su FRF e almeno un metodo BAD su PSD. Con una sola famiglia: entrambi i metodi di quella famiglia sono BAD."),
            ("Coh: ...", "Indicatore separato di qualità misura; non cambia il verdetto GOOD/SUSPECT/BAD."),
            ("Overload / Averages", "Campi OCR o contatore averages, utili per contesto e logging."),
        ],
        widths=[1.35, 4.95],
    )
    add_bullets(
        doc,
        [
            "Sample Freq (Hz) imposta la frequenza di polling; internamente screenshot_interval = 1 / frequenza.",
            "Events Only salva/logga solo quando cambiano i risultati rilevanti; disattivarlo produce logging continuo circa ogni secondo.",
            "Overlay mostra le ROI a schermo per controllare rapidamente l'allineamento.",
            "Audio è disponibile solo se sounddevice è installato e il dispositivo audio è accessibile.",
        ],
    )

    add_h(doc, "Log, immagini ed export UNV", 2)
    add_table(
        doc,
        ["Opzione", "Risultato"],
        [
            ("Verbose Log: Config/Mask/OCR/FFT/Lowpass/Classify/FileSave", "Aggiunge dettagli diagnostici al log testuale."),
            ("Image Logs: ROI", "Salva la porzione catturata per ogni hit."),
            ("Image Logs: Masks", "Salva la maschera binaria HSV."),
            ("Image Logs: Signal", "Salva il segnale ricostruito in unità fisiche."),
            ("Image Logs: FFT", "Salva spettro, cutoff e rapporto energia alta frequenza."),
            ("Image Logs: Lowpass", "Salva confronto segnale originale/filtro Lowpass."),
            ("Image Logs: Residual", "Salva residuo, soglia dinamica ed exceedances."),
            ("Image Logs: Summary", "Salva riepilogo della run."),
            (".unv", "Esporta Dataset 58 in signal_logs con parte immaginaria zero."),
        ],
        widths=[2.45, 3.85],
    )
    if "real_roi" in assets:
        add_image(doc, assets["real_roi"], "Figura 5 - Esempio reale di ROI salvata da USMA nei log immagine.", width=4.6)
    for key, caption in [
        ("real_signal", "Figura 6 - Log reale: segnale ricostruito in unità fisiche."),
        ("real_fft", "Figura 7 - Log reale: spettro FFT con cutoff e rapporto energia."),
        ("real_lowpass", "Figura 8 - Log reale: confronto segnale originale e Lowpass."),
        ("real_residual", "Figura 9 - Log reale: residuo, soglia dinamica ed exceedances."),
    ]:
        if key in assets:
            add_image(doc, assets[key], caption, width=6.35)

    add_h(doc, "Troubleshooting operativo", 2)
    add_table(
        doc,
        ["Sintomo", "Controllo consigliato"],
        [
            ("L'app non parte", "Apri run_log.txt e verifica che la cartella python/ sia presente."),
            ("Tesseract non trovato", "Verifica external/tesseract/ e le dipendenze OCR."),
            ("La maschera HSV non viene salvata", "Premi Apply nella finestra HSV e poi Save Config nello strumento ROI."),
            ("Molti falsi BAD", "Rivedi HSV, poi calibra; in emergenza alza fft_energy_ratio_threshold o fft_cutoff_frequency."),
            ("LP non segnala mai", "Abbassa exceedance_ratio_threshold o relative_residual_ratio, dopo aver controllato HSV."),
            ("La memoria cresce in sessioni lunghe", "La history GUI è limitata a 25 hit; per sessioni molto lunghe riavvia periodicamente."),
            ("UNV non contiene dati attesi", "Verifica X/Y axis scaling, unità ROI, punti Hammer/Response e opzione .unv abilitata."),
        ],
        widths=[2.05, 4.25],
    )
    doc.add_page_break()


def add_theory_wiki(doc: Document, assets: dict[str, Path]) -> None:
    add_h(doc, "Parte 2 - Wiki teorica", 1)
    add_callout(
        doc,
        "Copertura della wiki",
        "Questa parte integra i contenuti teorici presenti in THEORY.md e nelle pagine usma/theory/pages: ricostruzione segnale, FFT, Lowpass, "
        "parametri, tuning, coerenza, classificazione, calibrazione, HSV, grafici diagnostici e tabella dei default.",
        fill="EAFBF6",
    )

    add_h(doc, "1. Ricostruzione del segnale e calibrazione asse Y", 2)
    add_image(doc, assets["reconstruction"], "Figura 10 - Ricostruzione 1-D e conversione fisica del segnale.", width=6.25)
    add_para(
        doc,
        "Ogni analisi parte da una ROI e da una maschera HSV. Per ogni colonna x, USMA calcola la riga media dei pixel bianchi "
        "della maschera. Le colonne senza pixel vengono riempite interpolando dalle colonne vicine. Poiché nelle immagini la riga 0 è in alto, "
        "il segnale viene invertito rispetto all'altezza della ROI.",
    )
    add_formula(doc, "signal_pixels[x] = height - mean_row_of_white_pixels[x]")
    add_formula(doc, "signal_physical[i] = y_min + (signal_pixels[i] / height) * (y_max - y_min)")
    add_bullets(
        doc,
        [
            "Pixel-space e physical-space non sono la stessa cosa: il secondo dipende dalla scala Y configurata.",
            "FFT e Lowpass operano su signal_pixels, quindi la classificazione non dipende da errori di scala fisica.",
            "Il segnale fisico resta necessario per display, interpretazione ed export UNV.",
            "USMA assume asse Y lineare. Scale log/dB sono metadata utili all'utente, ma la ricostruzione non corregge automaticamente assi non lineari.",
            "La qualità della maschera è validata: copertura pixel tra circa 0.0005 e 0.4 e continuità colonne almeno 0.15.",
        ],
    )

    add_h(doc, "2. Metodo FFT", 2)
    add_para(
        doc,
        "Il metodo FFT misura quanta energia della forma del segnale cade sopra una frequenza spaziale normalizzata. "
        "La frequenza non è in Hz temporali: il segnale è campionato orizzontalmente in pixel, quindi l'asse FFT va da 0 a 0.5, con 0.5 pari a Nyquist.",
    )
    add_numbers(
        doc,
        [
            "Sottrae la media al segnale pixel.",
            "Calcola rfft e spettro di magnitudine.",
            "Costruisce l'asse rfftfreq(N, 1).",
            "Somma l'energia sopra fft_cutoff_frequency.",
            "Calcola energy_ratio = high_freq_energy / total_energy.",
            "Vota BAD se energy_ratio > fft_energy_ratio_threshold.",
        ],
    )
    add_formula(doc, "energy_ratio = sum(|FFT[f >= cutoff]|^2) / sum(|FFT|^2)")
    add_table(
        doc,
        ["Parametro", "Default", "Effetto se lo aumenti", "Effetto se lo abbassi"],
        [
            ("fft_cutoff_frequency", "0.07", "Meno spettro è considerato alta frequenza; FFT diventa meno sensibile.", "Più spettro è alta frequenza; FFT diventa più sensibile."),
            ("fft_energy_ratio_threshold", "0.006", "Serve più energia HF per BAD; meno falsi positivi.", "Basta meno energia HF; più sensibilità ai colpi sporchi."),
        ],
        widths=[1.6, 0.8, 2.0, 1.9],
    )

    add_h(doc, "3. Metodo Lowpass", 2)
    add_para(
        doc,
        "Il metodo Lowpass separa la componente liscia dalla componente rapida del segnale pixel. Dopo un Butterworth lowpass applicato con sosfiltfilt, "
        "USMA calcola il residuo e conta quanti campioni superano una soglia dinamica relativa al massimo residuo assoluto.",
    )
    add_formula(doc, "residual = (signal_pixels - mean) - lowpass(signal_pixels - mean)")
    add_formula(doc, "dynamic_threshold = relative_residual_ratio * max(abs(residual))")
    add_formula(doc, "exceedance_ratio = count(abs(residual) > dynamic_threshold) / N")
    add_bullets(
        doc,
        [
            "Voto Lowpass BAD se exceedance_ratio > exceedance_ratio_threshold.",
            "La soglia è relativa e adimensionale: non dipende dall'unità fisica della ROI.",
            "Se il residuo è praticamente zero, la soglia dinamica diventa zero e il segnale perfettamente liscio resta GOOD.",
            "Nel codice attuale i campi filtered_physical e residual_physical conservano dati pixel-space per retrocompatibilità dei nomi.",
        ],
    )
    add_table(
        doc,
        ["Parametro", "Default", "Ruolo"],
        [
            ("lowpass_cutoff", "0.07", "Cutoff normalizzato del filtro; controlla cosa viene considerato inviluppo liscio."),
            ("lowpass_filter_order", "7", "Ordine Butterworth; con sosfiltfilt l'ordine effettivo è doppio. È una manopola a basso impatto."),
            ("relative_residual_ratio", "0.10", "Frazione del massimo residuo che definisce la soglia dinamica."),
            ("exceedance_ratio_threshold", "0.7", "Frazione minima di campioni oltre soglia per votare BAD."),
        ],
        widths=[1.8, 0.8, 3.7],
    )
    add_table(
        doc,
        ["Order", "Ordine effettivo", "Attenuazione a 2x cutoff"],
        [
            ("3", "6", "~ -72 dB"),
            ("5", "10", "~ -120 dB"),
            ("7", "14", "~ -168 dB"),
            ("9", "18", "~ -216 dB"),
        ],
        widths=[1.4, 2.3, 2.6],
    )
    add_callout(
        doc,
        "Alternativa non implementata",
        "La wiki discute anche una soglia relativa alla media del residuo invece che al massimo. USMA usa l'opzione max-based perché separa meglio "
        "molte popolazioni Good/Bad tipiche degli impact test.",
    )

    add_h(doc, "4. Guida pratica al tuning", 2)
    add_table(
        doc,
        ["Obiettivo", "Intervento"],
        [
            ("FFT troppo sensibile", "Alza fft_energy_ratio_threshold oppure alza fft_cutoff_frequency."),
            ("FFT non segnala colpi evidentemente sporchi", "Abbassa fft_energy_ratio_threshold oppure abbassa fft_cutoff_frequency."),
            ("Lowpass non si attiva mai", "Abbassa exceedance_ratio_threshold o relative_residual_ratio."),
            ("Lowpass segnala colpi puliti", "Alza relative_residual_ratio o exceedance_ratio_threshold."),
            ("Maschera rumorosa o bucata", "Rifai la calibrazione HSV prima di ritoccare parametri numerici."),
            ("Soglie poco convincenti", "Svuota la calibrazione e ricampiona segnali variati; obiettivo Level 3/4."),
            ("Coerenza segnala ma hit sembra buono", "È normale: coerenza è canale qualità misura, non causa BAD verdict."),
        ],
        widths=[2.6, 3.7],
    )
    add_callout(
        doc,
        "Insight principale",
        "HSV governa ciò che entra nel segnale. Se la maschera è sbagliata, nessuna soglia FFT o Lowpass può essere davvero affidabile.",
        fill="FFF7E6",
    )

    add_h(doc, "5. Analisi di coerenza", 2)
    add_image(doc, assets["coherence"], "Figura 11 - Badness di coerenza: area normalizzata sotto 1 - gamma^2.", width=6.25)
    add_para(
        doc,
        "La coerenza viene ricostruita con lo stesso approccio HSV, ma il segnale fisico è forzato nell'intervallo [0, 1]. "
        "USMA integra 1 - gamma^2 sulla banda di frequenza visibile e normalizza per l'ampiezza della banda.",
    )
    add_formula(doc, "B = integral(1 - gamma^2) df / delta_f")
    add_table(
        doc,
        ["Metrica", "Formula", "Flag default"],
        [
            ("Mean coherence", "mean(signal_physical)", "-"),
            ("Min coherence", "min(signal_physical)", "-"),
            ("Normalised badness B", "integral(1 - gamma^2) df / delta_f", "> 0.30"),
            ("Per-band badness B_k", "integral_band(1 - gamma^2) / delta_f_band", "-"),
            ("Degradation", "(B - B_run_start) / B_run_start", "> 0.20"),
        ],
        widths=[1.75, 3.25, 1.3],
    )
    add_bullets(
        doc,
        [
            "USMA divide la curva in quattro bande e calcola una badness per ciascuna.",
            "Il trend run-level può essere IMPROVING, STABLE, DEGRADING o INSUFFICIENT_DATA.",
            "Coerenza non entra nella classificazione hit: è separata per non confondere qualità del setup con qualità del singolo colpo.",
        ],
    )

    add_h(doc, "6. Classificazione hit", 2)
    add_para(
        doc,
        "La classificazione aggrega fino a quattro voti: FRF-FFT, FRF-LP, PSD-FFT e PSD-LP. "
        "Ogni metodo produce BAD o non-BAD; il classificatore decide il testo finale e il colore.",
    )
    add_formula(doc, "not any_bad -> GOOD HIT")
    add_formula(doc, "all_bad     -> BAD HIT (dettaglio metodi)")
    add_formula(doc, "otherwise   -> SUSPECT (dettaglio metodi)")
    add_bullets(
        doc,
        [
            "Con FRF e PSD presenti, BAD richiede almeno un voto BAD su FRF e almeno un voto BAD su PSD.",
            "Con una sola famiglia presente, BAD richiede entrambi i metodi BAD in quella famiglia.",
            "Un singolo metodo rosso non basta a bocciare l'hit: produce SUSPECT per ridurre falsi positivi.",
            "Il dettaglio, ad esempio FRF-FFT+PSD-LP, indica quali metodi hanno fallito.",
        ],
    )

    add_h(doc, "7. Motore di calibrazione", 2)
    add_para(
        doc,
        "Il motore HybridCalibrationEngine riceve segnali etichettati dall'utente come GOOD o BAD e stima parametri separati per FRF e PSD quando "
        "ci sono abbastanza esempi per la famiglia di segnale. Il requisito minimo operativo è 3 GOOD e 3 BAD.",
    )
    add_table(
        doc,
        ["Stimatore", "Quando entra", "Cosa fa"],
        [
            ("PercentileBoundaryEstimator", "Level 1+", "Separa distribuzioni Good/Bad con midpoint se separate o 95° percentile dei Good se sovrapposte; fa sweep di cutoff e ratio."),
            ("BayesianThresholdEstimator", "Level 2+", "Usa griglie uniformi, likelihood sigmoidale, stima MAP e intervalli credibili al 95%."),
            ("ROCYoudenEstimator", "Level 3+", "Sceglie soglie che massimizzano Youden J = TPR - FPR e calcola AUC."),
        ],
        widths=[2.0, 1.1, 3.2],
    )
    add_table(
        doc,
        ["Parametro Bayes", "Range", "Punti griglia"],
        [
            ("fft_energy_ratio_threshold", "0.0005 - 0.15", "300"),
            ("exceedance_ratio_threshold", "0.05 - 0.99", "300"),
            ("relative_residual_ratio", "0.02 - 0.50", "300"),
        ],
        widths=[2.4, 2.0, 1.9],
    )
    add_table(
        doc,
        ["Livello", "Segnali", "Merge"],
        [
            ("0", "< 6 o minimo non raggiunto", "Default; non calibrato."),
            ("1", "6-7", "Percentile puro."),
            ("2", "8-11", "0.4 * percentile + 0.6 * Bayesian."),
            ("3", "12-15", "Media dei tre stimatori se concordano entro 20%; altrimenti scarta l'outlier."),
            ("4", "16+", "Come Level 3, ma concordanza più stretta: 15%."),
        ],
        widths=[0.75, 2.1, 3.45],
    )
    add_bullets(
        doc,
        [
            "I livelli sono ordinali, non probabilità assolute.",
            "La calibrazione non verifica automaticamente se GOOD e BAD sono davvero separabili.",
            "Non usa un modello multivariato: ogni metrica è calibrata indipendentemente.",
            "Non fa online learning incrementale; ricalcola dagli esempi salvati.",
            "Non ritocca continuamente fft_cutoff_frequency o lowpass_cutoff dopo la prima stima, salvo lo sweep negli stimatori previsti.",
        ],
    )

    add_h(doc, "8. Calibrazione colore HSV", 2)
    add_para(
        doc,
        "La calibrazione HSV è una soglia per canale con cv2.inRange. Un pixel è incluso se e solo se rispetta contemporaneamente i limiti H, S e V.",
    )
    add_formula(doc, "mask = cv2.inRange(hsv, hsv_lower, hsv_upper)")
    add_table(
        doc,
        ["Canale", "Range OpenCV", "Nota"],
        [
            ("Hue H", "0 - 179", "Non 0 - 360; attenzione ai picker esterni."),
            ("Saturation S", "0 - 255", "Controlla purezza del colore."),
            ("Value V", "0 - 255", "Default upper 240 per escludere bianco puro."),
        ],
        widths=[1.3, 1.6, 3.4],
    )
    add_bullets(
        doc,
        [
            "USMA non usa auto-clustering, k-means o Otsu: la regolazione è interattiva.",
            "La preview mostra originale, maschera e immagine filtrata.",
            "Una maschera buona include la traccia anche nei punti anti-aliasing ma scarta sfondo, griglia e testo.",
        ],
    )

    add_h(doc, "9. Grafici diagnostici", 2)
    add_table(
        doc,
        ["Grafico", "Cosa visualizza", "Uso"],
        [
            ("Signal plot", "signal_physical contro indice/frequenza ROI.", "Controllare ampiezza e forma esportabile."),
            ("FFT plot", "Magnitudine FFT, cutoff, area alta frequenza e energy_ratio.", "Capire perché FFT vota BAD."),
            ("Lowpass comparison", "Segnale pixel originale e versione filtrata.", "Visualizzare cosa viene considerato inviluppo liscio."),
            ("Residual analysis", "Residuo pixel, soglia dinamica, campioni oltre soglia.", "Capire il voto Lowpass."),
            ("Run summary", "Distribuzioni/count degli hit nella run.", "Controllare tendenze e outlier."),
            ("Calibration distributions", "Istogrammi GOOD/BAD in densità, soglia e intervalli credibili.", "Valutare separazione e incertezza."),
            ("ROC subplot", "Curve FPR/TPR, AUC e punto Youden.", "Stimare discriminatività di ogni metrica."),
            ("Convergence", "Soglie stimate al crescere dei campioni.", "Vedere se la calibrazione si stabilizza."),
        ],
        widths=[1.55, 2.9, 1.85],
    )
    add_callout(
        doc,
        "Densità negli istogrammi",
        "Gli istogrammi di calibrazione usano density=True: l'area totale di ogni popolazione vale 1. Questo rende confrontabili GOOD e BAD anche se hanno numerosità diverse.",
    )

    add_h(doc, "10. Parametri default", 2)
    add_default_parameter_tables(doc)

    add_h(doc, "11. Cosa resta fuori dalla teoria", 2)
    add_bullets(
        doc,
        [
            "OCR: legge testo e metadati, ma non trasforma il segnale numerico.",
            "UNV exporter: serializza Dataset 58 senza modificare la ricostruzione, salvo formattazione e parte immaginaria zero.",
            "Image logger: salva immagini diagnostiche, non cambia l'analisi.",
            "Monitor loop: pianifica catture e dispatch eventi; la matematica è nei moduli di analisi.",
        ],
    )

    add_h(doc, "12. Changelog teorico LP v0.10.2+", 2)
    add_bullets(
        doc,
        [
            "Il metodo Lowpass è passato da signal_physical a signal_pixels.",
            "La soglia residua assoluta è stata sostituita da relative_residual_ratio.",
            "Il filtro usa butter(output='sos') più sosfiltfilt per stabilità numerica.",
            "Entrambi i metodi principali, FFT e LP, operano ora in pixel-space; la scala fisica resta per visualizzazione ed export.",
        ],
    )
    doc.add_page_break()


def add_default_parameter_tables(doc: Document) -> None:
    add_table(
        doc,
        ["Parametro", "Default", "Range", "Effetto"],
        [
            ("fft_cutoff_frequency", "0.07", "0.02 - 0.25", "Confine tra bassa e alta frequenza spaziale."),
            ("fft_energy_ratio_threshold", "0.006", "0.0001 - 0.1", "Frazione minima di energia HF per voto FFT BAD."),
            ("lowpass_cutoff", "0.07", "0.03 - 0.2", "Cutoff Butterworth normalizzato."),
            ("lowpass_filter_order", "7", "-", "Ordine Butterworth, effettivo 14 con sosfiltfilt."),
            ("relative_residual_ratio", "0.10", "0.02 - 0.50", "Frazione del massimo residuo per la soglia dinamica."),
            ("exceedance_ratio_threshold", "0.7", "0.2 - 0.99", "Frazione minima di exceedance per voto LP BAD."),
        ],
        widths=[2.0, 0.8, 1.35, 2.15],
    )
    add_para(doc, "PSD usa un set parallelo con gli stessi default e prefisso psd_.")
    add_table(
        doc,
        ["Parametro", "Default", "Range", "Effetto"],
        [
            ("coherence_threshold", "0.30", "0 - 1", "Massima badness normalizzata prima dell'avviso."),
            ("coherence_degradation_pct", "0.20", "0 - 1", "Aumento relativo nella run che segnala degrado."),
            ("hsv_lower", "[0, 0, 0]", "per canale", "Limite inferiore HSV."),
            ("hsv_upper", "[179, 255, 240]", "per canale", "Limite superiore HSV, V sotto 255 per escludere bianco puro."),
            ("screenshot_interval", "0.25 s", "-", "Periodo loop monitor; equivale a 4 Hz."),
            ("hits_per_run", "5", "-", "Numero atteso di hit per run."),
            ("monitor_index", "1", "-", "Monitor fisico catturato da mss."),
        ],
        widths=[2.05, 1.05, 1.15, 2.05],
    )


def add_references(doc: Document) -> None:
    add_h(doc, "Parte 3 - Fonti tecniche e controllo qualità", 1)
    add_h(doc, "Fonti usate", 2)
    add_table(
        doc,
        ["Fonte", "Contenuto usato"],
        [
            ("README.md", "Quick start, struttura cartelle, tipi ROI, classificazione, export e troubleshooting."),
            ("THEORY.md", "Riferimento matematico completo."),
            ("usma/theory/pages/*.txt", "Pagine wiki integrate nell'applicazione."),
            ("usma/models.py", "Versione, dataclass, default parametri e motore calibrazione."),
            ("usma/analysis/signal.py", "Ricostruzione, FFT, Lowpass, soglia dinamica ed exceedance."),
            ("usma/analysis/coherence.py", "Badness di coerenza e tracking per banda."),
            ("usma/analysis/classifier.py", "Logica esatta GOOD/SUSPECT/BAD."),
            ("usma/gui/*.py", "Controlli GUI, strumenti ROI, HSV, graph viewer e calibrazione live."),
            ("usma/export/unv.py", "Serializzazione UNV Dataset 58."),
            ("image_logs/", "Esempi reali di grafici diagnostici."),
        ],
        widths=[2.15, 4.15],
    )
    add_h(doc, "Pagine wiki coperte", 2)
    add_table(
        doc,
        ["Pagina", "Dove viene integrata"],
        [
            ("signal_reconstruction", "Ricostruzione segnale e calibrazione asse Y."),
            ("fft_method, fft_cutoff_frequency, fft_energy_ratio_threshold", "Metodo FFT e parametri."),
            ("lowpass_method, lowpass_cutoff, lowpass_filter_order", "Metodo Lowpass e filtro."),
            ("residual_threshold, exceedance_ratio_threshold", "Soglia residua relativa ed exceedance."),
            ("practical_tuning", "Guida pratica al tuning."),
            ("coherence_analysis", "Analisi di coerenza."),
            ("hit_classification", "Classificazione hit."),
            ("calibration_engine", "Motore di calibrazione ibrido."),
            ("hsv_calibration", "Filtro HSV."),
            ("default_parameters", "Tabelle di riferimento parametri."),
        ],
        widths=[2.8, 3.5],
    )
    add_h(doc, "Controlli di accuratezza applicati", 2)
    add_bullets(
        doc,
        [
            f"Versione letta dal codice: USMA v{APP_VERSION}.",
            "Valori numerici default verificati in usma/models.py e THEORY.md.",
            "Logica classificazione verificata in usma/analysis/classifier.py.",
            "Grafici sintetici generati da formule esplicitate nella wiki; grafici scuri importati da image_logs come esempi reali.",
            "Nessuna schermata UI è stata inventata: le figure operative sono diagrammi o log già presenti.",
        ],
    )


def build_document() -> Path:
    assets = generate_assets()
    doc = Document()
    configure_styles(doc)
    add_cover(doc, assets)
    add_index(doc)
    add_operational_guide(doc, assets)
    add_theory_wiki(doc, assets)
    add_references(doc)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
